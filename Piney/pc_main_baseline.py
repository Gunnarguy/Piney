import os
import argparse
import openai
from pinecone import Pinecone, ServerlessSpec
import tiktoken  # Ensure you have tiktoken installed (pip install tiktoken)
import sys

from dotenv import load_dotenv

load_dotenv()  # Loads variables from .env into os.environ

# Optional libraries for file processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx  # for processing .docx files
except ImportError:
    docx = None

# -------------------- CONFIGURATION --------------------
def load_config():
    pinecone_key = os.getenv("PINECONE_API_KEY", "")
    openai_key = os.getenv("OPENAI_API_KEY", "")
    
    if not pinecone_key or not openai_key:
        raise EnvironmentError("Please set both PINECONE_API_KEY and OPENAI_API_KEY environment variables.")
    
    # Check if the keys are valid (optional, but recommended)
    # You might want to ping the OpenAI and Pinecone APIs to validate the keys
    # This example skips the validation for brevity
    
    return pinecone_key, openai_key

# -------------------- ARGUMENT PARSING --------------------
def parse_args():
    parser = argparse.ArgumentParser(description="Run Pinecone Embedding Pipeline with interactive or non-interactive mode")
    parser.add_argument("--index", type=str, help="Name of the Pinecone index to use or create")
    parser.add_argument("--namespace", type=str, help="Namespace to use (if omitted, will prompt or use default)")
    parser.add_argument("--directory", type=str, help="Path to the directory containing files")
    parser.add_argument("--non-interactive", action="store_true", help="Run in non-interactive mode using provided defaults")
    return parser.parse_args()

# -------------------- INTERACTIVE INDEX SELECTION --------------------
def select_or_create_index(pc, provided_index=None, non_interactive=False):
    existing_indexes = list(pc.list_indexes().names())
    if provided_index:
        if provided_index in existing_indexes:
            print(f"Using provided existing index: {provided_index}")
            return provided_index
        else:
            print(f"Provided index '{provided_index}' does not exist. Creating it...")
            pc.create_index(
                name=provided_index,
                dimension=3072,  # Adjust based on OpenAI model used
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
            return provided_index

    if non_interactive:
        # Fallback default in non-interactive mode
        default_index = "default_index"
        if default_index not in existing_indexes:
            print(f"Creating default index '{default_index}'...")
            pc.create_index(
                name=default_index,
                dimension=3072,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
        else:
            print(f"Using existing default index: {default_index}")
        return default_index

    # Interactive mode
    print("Existing indexes:", existing_indexes)
    while True:
        choice = input("Do you want to (s)elect an existing index or (c)reate a new one? [s/c]: ").strip().lower()
        if choice not in ("s", "c"):
            print("Invalid choice. Please enter 's' to select or 'c' to create.")
            continue
        if choice == "s":
            index_name = input("Enter the name of the existing index: ").strip()
            if index_name in existing_indexes:
                return index_name
            else:
                create_choice = input(f"Index '{index_name}' does not exist. Would you like to create it instead? (y/n): ").strip().lower()
                if create_choice == "y":
                    print(f"Creating new index '{index_name}'...")
                    pc.create_index(
                        name=index_name,
                        dimension=3072,
                        metric="cosine",
                        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                    )
                    return index_name
                else:
                    print("Let's try again.")
                    continue
        elif choice == "c":
            index_name = input("Enter a name for the new index: ").strip()
            if index_name in existing_indexes:
                print(f"Index '{index_name}' already exists. Using the existing index.")
            else:
                print(f"Creating new index '{index_name}'...")
                pc.create_index(
                    name=index_name,
                    dimension=3072,
                    metric="cosine",
                    spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                )
            return index_name

# -------------------- INTERACTIVE NAMESPACE SELECTION --------------------
def select_or_create_namespace(provided_namespace=None, non_interactive=False):
    if provided_namespace is not None:
        print(f"Using provided namespace: '{provided_namespace}'")
        return provided_namespace
    if non_interactive:
        default_namespace = "default_namespace"
        print(f"Non-interactive mode: using default namespace '{default_namespace}'")
        return default_namespace

    namespace = input("Enter a namespace to use (or press Enter to use the default namespace): ").strip()
    return namespace if namespace != "" else ""

# -------------------- DOCUMENT CHUNKING --------------------
def chunk_text(text, max_tokens=1000):
    """
    Splits text into chunks of approximately max_tokens tokens each.
    Uses tiktoken for tokenization and ensures chunks do not exceed the limit.
    """
    enc = tiktoken.get_encoding("cl100k_base")
    tokens = enc.encode(text)
    
    # If the text is within the limit, return it as a single chunk.
    if len(tokens) <= max_tokens:
        return [text]
    
    chunks = []
    start = 0
    while start < len(tokens):
        end = start + max_tokens
        chunk_tokens = tokens[start:end]
        chunk_text = enc.decode(chunk_tokens)
        chunks.append(chunk_text)
        start = end
    return chunks

# -------------------- FILE PROCESSING --------------------
def process_file(filepath):
    """
    Processes a file and extracts text content.
    Supports .txt, .md, .py, .json, .csv, .pdf, and .docx files.
    """
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    try:
        if ext in ['.txt', '.md', '.py', '.json', '.csv']:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext == '.pdf' and PyPDF2:
            with open(filepath, "rb") as f:
                try:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() or ""
                except Exception as e:
                    print(f"Error reading PDF {filepath}: {e}")
                    return []  # Skip this file if PDF reading fails
        elif ext in ['.docx'] and docx:
            try:
                doc = docx.Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                print(f"Error reading DOCX {filepath}: {e}")
                return []  # Skip if DOCX fails
        else:
            # For unknown types, attempt to read as text
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except IOError as e:
        print(f"IOError processing file {filepath}: {e}")
        return []  # Return an empty list if there's an IOError
    except Exception as e:
        print(f"Unexpected error processing file {filepath}: {e}")
        return []

    # Chunk the text if it's too long.
    chunks = chunk_text(text)
    print(f"Processed '{filepath}' into {len(chunks)} chunk(s).")
    return chunks

def process_directory(directory):
    documents = []
    doc_id = 1
    for root, dirs, files in os.walk(directory):
        for file in files:
            filepath = os.path.join(root, file)
            chunks = process_file(filepath)
            for chunk in chunks:
                if chunk.strip():
                    documents.append({
                        "id": f"doc{doc_id}",
                        "text": chunk
                    })
                    doc_id += 1
    return documents

# -------------------- CONFIGURATION --------------------
def load_config():
    pinecone_key = os.getenv("PINECONE_API_KEY", "")
    openai_key = os.getenv("OPENAI_API_KEY", "")
    
    if not pinecone_key or not openai_key:
        raise EnvironmentError("Please set both PINECONE_API_KEY and OPENAI_API_KEY environment variables.")
    
    # Check if the keys are valid (optional, but recommended)
    # You might want to ping the OpenAI and Pinecone APIs to validate the keys
    # This example skips the validation for brevity
    
    return pinecone_key, openai_key

# -------------------- ARGUMENT PARSING --------------------
def parse_args():
    parser = argparse.ArgumentParser(description="Run Pinecone Embedding Pipeline with interactive or non-interactive mode")
    parser.add_argument("--index", type=str, help="Name of the Pinecone index to use or create")
    parser.add_argument("--namespace", type=str, help="Namespace to use (if omitted, will prompt or use default)")
    parser.add_argument("--directory", type=str, help="Path to the directory containing files")
    parser.add_argument("--non-interactive", action="store_true", help="Run in non-interactive mode using provided defaults")
    return parser.parse_args()

# -------------------- INTERACTIVE INDEX SELECTION --------------------
def select_or_create_index(pc, provided_index=None, non_interactive=False):
    existing_indexes = list(pc.list_indexes().names())
    if provided_index:
        if provided_index in existing_indexes:
            print(f"Using provided existing index: {provided_index}")
            return provided_index
        else:
            print(f"Provided index '{provided_index}' does not exist. Creating it...")
            pc.create_index(
                name=provided_index,
                dimension=3072,  # Adjust based on OpenAI model used
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
            return provided_index

    if non_interactive:
        # Fallback default in non-interactive mode
        default_index = "default_index"
        if default_index not in existing_indexes:
            print(f"Creating default index '{default_index}'...")
            pc.create_index(
                name=default_index,
                dimension=3072,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
        else:
            print(f"Using existing default index: {default_index}")
        return default_index

    # Interactive mode
    print("Existing indexes:", existing_indexes)
    while True:
        choice = input("Do you want to (s)elect an existing index or (c)reate a new one? [s/c]: ").strip().lower()
        if choice not in ("s", "c"):
            print("Invalid choice. Please enter 's' to select or 'c' to create.")
            continue
        if choice == "s":
            index_name = input("Enter the name of the existing index: ").strip()
            if index_name in existing_indexes:
                return index_name
            else:
                create_choice = input(f"Index '{index_name}' does not exist. Would you like to create it instead? (y/n): ").strip().lower()
                if create_choice == "y":
                    print(f"Creating new index '{index_name}'...")
                    pc.create_index(
                        name=index_name,
                        dimension=3072,
                        metric="cosine",
                        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                    )
                    return index_name
                else:
                    print("Let's try again.")
                    continue
        elif choice == "c":
            index_name = input("Enter a name for the new index: ").strip()
            if index_name in existing_indexes:
                print(f"Index '{index_name}' already exists. Using the existing index.")
            else:
                print(f"Creating new index '{index_name}'...")
                pc.create_index(
                    name=index_name,
                    dimension=3072,
                    metric="cosine",
                    spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                )
            return index_name

# -------------------- INTERACTIVE NAMESPACE SELECTION --------------------
def select_or_create_namespace(provided_namespace=None, non_interactive=False):
    if provided_namespace is not None:
        print(f"Using provided namespace: '{provided_namespace}'")
        return provided_namespace
    if non_interactive:
        default_namespace = "default_namespace"
        print(f"Non-interactive mode: using default namespace '{default_namespace}'")
        return default_namespace

    namespace = input("Enter a namespace to use (or press Enter to use the default namespace): ").strip()
    return namespace if namespace != "" else None

# -------------------- DOCUMENT CHUNKING --------------------
def chunk_text(text, max_tokens=1000):
    """
    Splits text into chunks of approximately max_tokens tokens each.
    Uses tiktoken for tokenization and ensures chunks do not exceed the limit.
    """
    enc = tiktoken.get_encoding("cl100k_base")
    tokens = tokenize_text(text)
    
    # If the text is within the limit, return it as a single chunk.
    if len(tokens) <= max_tokens:
        return [text]
    
    chunks = []
    start = 0
    while start < len(tokens):
        end = start + max_tokens
        chunk_tokens = tokens[start:end]
        chunk_text = enc.decode(chunk_tokens)
        chunks.append(chunk_text)
        start = end
    return chunks

def process_file(filepath):
    """
    Processes a file and extracts text content.
    Supports .txt, .md, .py, .json, .csv, .pdf, and .docx files.
    """
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    try:
        if ext in ['.txt', '.md', '.py', '.json', '.csv']:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext == '.pdf' and PyPDF2:
            with open(filepath, "rb") as f:
                try:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() or ""
                except Exception as e:
                    print(f"Error reading PDF {filepath}: {e}")
                    return []  # Skip this file if PDF reading fails
        elif ext in ['.docx'] and docx:
            try:
                doc = docx.Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                print(f"Error reading DOCX {filepath}: {e}")
                return []  # Skip if DOCX fails
        else:
            # For unknown types, attempt to read as text
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except IOError as e:
        print(f"IOError processing file {filepath}: {e}")
        return []  # Return an empty list if there's an IOError
    except Exception as e:
        print(f"Unexpected error processing file {filepath}: {e}")
        return []

    # Chunk the text if it's too long.
    chunks = chunk_text(text)
    print(f"Processed '{filepath}' into {len(chunks)} chunk(s).")
    return chunks

def process_directory(directory):
    documents = []
    doc_id = 1
    for root, dirs, files in os.walk(directory):
        for file in files:
            filepath = os.path.join(root, file)
            chunks = process_file(filepath)
            for chunk in chunks:
                if chunk.strip():
                    documents.append({
                        "id": f"doc{doc_id}",
                        "text": chunk
                    })
                    doc_id += 1
    return documents

def initialize(args):
    """
    Initializes the OpenAI and Pinecone clients, selects the index and namespace.
    
    Args:
        args: The command-line arguments.
    
    Returns:
        A tuple containing the OpenAI client, Pinecone client, index, and namespace.
    """
    PINECONE_API_KEY, OPENAI_API_KEY = load_config()

    # Initialize OpenAI API client
    client = openai.OpenAI(api_key=OPENAI_API_KEY)

    # Initialize Pinecone Client (New API)
    pc = Pinecone(api_key=PINECONE_API_KEY)

    # Determine index name using provided argument or interactive/default mode
    index_name = select_or_create_index(pc, provided_index=args.index, non_interactive=args.non_interactive)
    index = pc.Index(index_name)

    # Determine namespace using provided argument or interactive/default mode
    namespace = select_or_create_namespace(provided_namespace=args.namespace, non_interactive=args.non_interactive)

    return client, pc, index, namespace


# -------------------- MAIN FUNCTION --------------------
def main():
    """
    Main function to run the Pinecone Embedding Pipeline.
    """
    args = parse_args()

    # -------------------- FUNCTION: EMBED TEXT --------------------
    def embed_text(texts, client):
        """
        Embeds the given texts using the OpenAI API.
        
        Args:
            texts: A list of texts to embed.
            client: The OpenAI client.
        
        Returns:
            A list of embeddings.
        """
        response = client.embeddings.create(
            model="text-embedding-3-large",
            input=texts
        )
        return [r.embedding for r in response.data]

    # -------------------- FUNCTION: UPSERT DATA WITH BATCHING --------------------
    def upsert_data(documents, index, namespace, client, args):
        """
        Upserts the given documents into the Pinecone index with batching.
        
        Args:
            documents: A list of documents to upsert.
            index: The Pinecone index.
            namespace: The namespace to use.
            client: The OpenAI client.
            args: The command-line arguments.
        """
        batch_size = args.batch_size if hasattr(args, 'batch_size') and args.batch_size else 100
        # Embed texts for all documents
        embeddings = embed_text([doc["text"] for doc in documents], client)
        vectors = [
            {
                "id": doc["id"],
                "values": emb,
                "metadata": {"text": doc["text"]}
            }
            for doc, emb in zip(documents, embeddings)
        ]
        ns = namespace if namespace else ""
        # Batch upsert vectors to avoid exceeding request size limits
        total_batches = (len(vectors) + batch_size - 1) // batch_size
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            index.upsert(batch, namespace=ns)
            print(f"Upserted batch {i // batch_size + 1} of {total_batches}")
        print(f"Successfully upserted {len(vectors)} documents into index '{args.index}'" +
              (f" under namespace '{ns}'." if ns else "."))

    # -------------------- FUNCTION: QUERY PINECONE --------------------
    def query_pinecone(query, index, namespace, client, args, top_k=3):
        """
        Queries the Pinecone index with the given query.
        
        Args:
            query: The query to use.
            index: The Pinecone index.
            namespace: The namespace to use.
            client: The OpenAI client.
            args: The command-line arguments.
            top_k: The number of results to return.
        
        Returns:
            The query results.
        """
        query_embedding = embed_text([query], client)[0]
        ns = namespace if namespace else ""
        results = index.query(
            namespace=ns,
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True
        )
        return results

    # Determine directory to process
    def get_directory(args):
        """
        Determines the directory to process based on the command-line arguments.
        
        Args:
            args: The command-line arguments.
        
        Returns:
            The directory to process.
        """
        if args.directory:
            directory = args.directory
        elif args.non_interactive:
            directory = "./documents"
            print(f"Non-interactive mode: using default directory '{directory}'")
        else:
            directory = input("Enter the path to the directory containing your files: ").strip()

        if not os.path.isdir(directory):
            print(f"Directory '{directory}' is invalid. Exiting.")
            sys.exit(1)

        # Check if the directory is readable
        if not os.access(directory, os.R_OK):
            print(f"Directory '{directory}' is not readable. Exiting.")
            sys.exit(1)

        return directory

    directory = get_directory(args)

    print("Processing files in directory. This may take a while...")
    documents = process_directory(directory)
    if not documents:
        print("No valid documents found. Exiting.")
        sys.exit(1)

    client, pc, index, namespace = initialize(args)
    upsert_data(documents, index, namespace, client, args)

    # Interactive query loop (only if not in non-interactive mode)
    if not args.non_interactive:
        while True:
            choice = input("Would you like to (q)uery the index or (e)xit? [q/e]: ").strip().lower()
            if choice == "q":
                user_query = input("Enter your query: ").strip()
                results = query_pinecone(user_query, index, namespace, client, args)
                print("\nTop Results:")
                for match in results["matches"]:
                    print(f"- {match['metadata']['text']} (Score: {match['score']:.4f})")
            elif choice == "e":
                print("Exiting. Goodbye!")
                break
            else:
                print("Invalid option. Please enter 'q' to query or 'e' to exit.")
    else:
        print("Non-interactive mode complete. Exiting.")


if __name__ == "__main__":
    main()
